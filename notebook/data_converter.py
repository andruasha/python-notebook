from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_COLOR_INDEX
from notebook.tag import map_tag_string, create_tag


def create_document_from_tag_dump(dump):
    """Create Document object from docx lib by tags data."""
    document = Document()
    paragraph = document.add_paragraph()
    document_data = parse_tag_dump_to_document_data(dump)
    for document_data_item in document_data:
        apply_document_data_item_on_document(paragraph, document_data_item)
    return document


def parse_tag_dump_to_document_data(dump):
    """Parse tags data to document data."""
    document_data = []
    current_text, current_tag = '', ''
    for tag_dump_item in dump:
        key, value, index = tag_dump_item
        if key == 'text':
            current_text = value
            document_data.append((current_text, current_tag))
        else:
            current_tag = value

    return document_data


def apply_document_data_item_on_document(document, document_data_item):
    """Apply single document data item on document."""
    text, tag = document_data_item
    family, size, weight, slant, underline, text_color, back_color = parse_and_convert_tag_items(
        tag)

    run = document.add_run(text)
    run.font.name = family
    run.font.size = size
    run.bold = weight
    run.italic = slant
    run.underline = underline
    run.font.color.rgb = text_color
    run.font.highlight_color = back_color


def parse_and_convert_tag_items(tag):
    """Parse tag data and convert it."""
    family, size, weight, slant, underline, text_color, back_color = map_tag_string(
        tag)
    return family, Pt(int(size)), weight == 'bold', slant == 'italic',\
        underline == '1', RGBColor.from_string(
            text_color[1:]), convert_color_to_highlight_color(back_color)


def convert_color_to_highlight_color(color):
    """Convert color type."""
    formatted_color = color[1:]
    if formatted_color == '000000':
        return WD_COLOR_INDEX.BLACK
    elif formatted_color == 'ffffff':
        return WD_COLOR_INDEX.WHITE
    elif formatted_color == 'ff0000':
        return WD_COLOR_INDEX.RED


def create_tag_dump_from_document_data(document):
    """Create tags data by document data."""
    tag_dump = []
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            text = run.text
            tag = parse_and_convert_run_to_tag(run)
            tag_dump.append((text, tag))
    return tag_dump


def parse_and_convert_run_to_tag(run):
    """Parse run and convert it to tag."""
    family = run.font.name
    size = int(run.font.size.pt)
    weight = 'bold' if run.bold else 'normal'
    slant = 'italic' if run.italic else 'roman'
    underline = '1' if run.underline else '0'
    text_color = '#' + str(run.font.color.rgb)
    back_color = convert_highlight_color_to_color(run.font.highlight_color)

    return create_tag(family, size, weight, slant,
                      underline, text_color, back_color)


def convert_highlight_color_to_color(highlight_color):
    """Convert color type."""
    if highlight_color == WD_COLOR_INDEX.BLACK:
        return '#000000'
    elif highlight_color == WD_COLOR_INDEX.WHITE:
        return '#ffffff'
    elif highlight_color == WD_COLOR_INDEX.RED:
        return '#ff0000'
